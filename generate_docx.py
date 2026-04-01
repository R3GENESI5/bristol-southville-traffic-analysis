"""
Generate SBLN Internal Working Document (DOCX)
Authors: Ali Bin Shahid & Matt Sanders
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = r"D:\Projects\Bristol"
CHARTS = os.path.join(BASE, "charts")
OUTPUT = os.path.join(BASE, "SBLN_Internal_Working_Document.docx")

doc = Document()

# ── Style tweaks ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

# ── Helper functions ──

def add_image(path, width=Inches(6), caption=None):
    """Add image with try/except for missing files."""
    full = os.path.join(CHARTS, path) if not os.path.isabs(path) else path
    try:
        doc.add_picture(full, width=width)
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].italic = True
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    except FileNotFoundError:
        doc.add_paragraph(f"[Image not found: {path}]")

def add_table(headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()  # spacer
    return table

def page_break():
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# 1. TITLE PAGE
# ═══════════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("South Bristol Liveable Neighbourhoods")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Internal Working Document")
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()

label = doc.add_paragraph()
label.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = label.add_run("CONFIDENTIAL — INTERNAL USE ONLY")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_paragraph()

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run("Ali Bin Shahid — Volunteer Data Analyst & Systems Engineer\nMatt Sanders — Community Advocate & Art Director")
run.font.size = Pt(13)

doc.add_paragraph()

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run("April 2026")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

page_break()

# ═══════════════════════════════════════════════════════════════════
# 2. PROBLEM STATEMENT
# ═══════════════════════════════════════════════════════════════════

doc.add_heading("Problem Statement", level=1)

doc.add_heading("Bristol City Council's Website Claim", level=2)
doc.add_paragraph(
    "Bristol City Council's website states:"
)
quote = doc.add_paragraph()
quote.style = doc.styles['Intense Quote'] if 'Intense Quote' in [s.name for s in doc.styles] else doc.styles['Quote']
quote.add_run(
    '"Traffic data show that high numbers of vehicles from outside the area use '
    'residential streets in Southville as a cut-through to other destinations."'
)

doc.add_paragraph(
    "This claim is used to justify the South Bristol Liveable Neighbourhoods (SBLN) programme, "
    "which proposes installing modal filters (road closures) to divide the Southville area into "
    "four zones."
)

doc.add_heading("The Question", level=2)
doc.add_paragraph(
    "Which specific traffic data proves this claim?"
)
doc.add_paragraph(
    "When community members asked Bristol City Council to identify which of the 94 survey "
    "spreadsheets support the website's claim, no specific answer was provided."
)

doc.add_heading("The Evidence Gap", level=2)
doc.add_paragraph(
    "To demonstrate that vehicles are \"cutting through\" residential streets, the survey data "
    "would need to show all four of the following:"
)
items = [
    "A vehicle enters the area from one boundary road",
    "Travels through the residential streets (not around the perimeter)",
    "Exits from another boundary road",
    "Does not stop at a destination within the area",
]
for i, item in enumerate(items, 1):
    doc.add_paragraph(f"{i}. {item}", style='List Number')

doc.add_paragraph(
    "Requirements 1 and 3 need ANPR cameras on the perimeter (which exist — Cameras 14 and 16). "
    "Requirements 2 and 4 need ANPR cameras within the Southville Zone — which do NOT exist."
)

doc.add_heading("FOI Request and Response", level=2)
doc.add_paragraph(
    "A Freedom of Information request was submitted in December 2024. It took 8 weeks to receive "
    "a response — double the statutory 4-week deadline. The eventual response was described by "
    'the requester as "a laughable misunderstanding of the data."'
)

doc.add_heading("WECA Funding Context", level=2)
doc.add_paragraph(
    "The SBLN programme received £1.32m in additional funding (on top of £0.6m already allocated) "
    "from WECA (West of England Combined Authority). This was approved by a slim 5:4 committee "
    "vote at the Transport and Connectivity Policy Committee in February 2025. "
    "WECA funding is only available for a liveable neighbourhood if it eliminates \"through traffic\" "
    "by installing modal filters."
)

page_break()

# ═══════════════════════════════════════════════════════════════════
# 3. MATT SANDERS' WECA STATEMENT
# ═══════════════════════════════════════════════════════════════════

doc.add_heading("Matt Sanders' WECA Statement", level=1)
doc.add_paragraph(
    "Statement to the West of England Combined Authority Joint Committee Meeting, 27 March 2026."
)
doc.add_paragraph(
    'Full title: "SBLN: The Fundamental Lie"'
).runs[0].italic = True

# Full statement as indented paragraphs
statement_sections = [
    ("1. The SBLN Modal Filters",
     "Bristol City Council is currently developing proposals for the South Bristol Liveable Neighbourhood.\n\n"
     "But funding from WECA (West of England Combined Authority) is ONLY available for a liveable neighbourhood, "
     "if it eliminates \"through traffic\" by installing modal filters.\n\n"
     "The design proposal for Southville does indeed consist of dividing this area into four \"zones\", "
     "separated by modal filters."),

    ("2. The Council Website Claim",
     "The Council Website says:\n\n"
     "\"Traffic data show that high numbers of vehicles from outside the area use residential streets in "
     "Southville as a cut-through to other destinations.\"\n\n"
     "But WHAT specific \"Traffic data\" actually PROVES this claim...?\n\n"
     "When we ask this question, we are directed to download \"Traffic Survey Results\" from June 2024 — "
     "94 complex spreadsheets, each with dozens of pages.\n\n"
     "However — when we ask WHICH specific spreadsheets support the website's claim — no answer is ever forthcoming."),

    ("3. ANPR Camera Sites",
     "The only way to demonstrate that vehicles are \"cutting through\" is to have ANPR cameras throughout "
     "the streets of Southville — and simultaneously — more cameras on all the perimeter roads.\n\n"
     "Six of the 94 spreadsheets relate to ANPR data. The camera location map reveals that — in June 2024 — "
     "there were 25 Cameras in South Bristol.\n\n"
     "But NONE of these cameras were WITHIN the \"Southville Zone.\"\n\n"
     "Though TWO cameras were on Southville's perimeter roads..."),

    ("4. Southville Camera Sites",
     "Camera 14 was on Coronation Road. Camera 16 was on North Street.\n\n"
     "The ANPR spreadsheets show that, on 11th June 2024, each of these cameras saw more than 1,000 vehicles, "
     "which were also detected by other cameras within the same 15 minutes.\n\n"
     "But only FOUR of those vehicles were detected by BOTH of these cameras.\n\n"
     "And there is no way of knowing if they drove around the PERIMETER of the Southville zone, or cut through "
     "the middle — as there were NO cameras WITHIN the zone.\n\n"
     "The vast majority of matching detections related to Camera 1 — at Clift House Road — as the vehicles "
     "seen by Camera 14 drove along Coronation Road — the NORTHERN PERIMETER — while those seen by Camera 16 "
     "drove along North Street — the SOUTHERN PERIMETER — all WITHOUT \"cutting through\" the Southville Zone."),

    ("5. The Fundamental Lie",
     "So, no — Bristol City Council has NO \"Traffic Data\", to prove this claim:\n\n"
     "\"Traffic data show that high numbers of vehicles from outside the area use residential streets in "
     "Southville as a cut-through to other destinations.\"\n\n"
     "Yet this false claim is being used to garner public support for the scheme.\n\n"
     "Based upon claims like this — at February's Transport and Connectivity Policy Committee, a slim 5:4 "
     "majority of Councillors were persuaded to approve a £1.32m increase to the then-current allocation of "
     "£0.6m, for development of the SBLN Full Business Case.\n\n"
     "But without coherent Traffic data, there is NO JUSTIFICATION for imposing modal filters, with the "
     "colossal disruption they would cause.\n\n"
     "The design has clearly been reverse-engineered — to \"solve\" a problem which does not really exist — "
     "in order to access the WECA funding package.\n\n"
     "So the entire South Bristol Liveable Neighbourhood scheme is nothing more than a house of cards — "
     "proposed on the basis of this Fundamental LIE.\n\n"
     "— Matt Sanders"),
]

for heading, text in statement_sections:
    p = doc.add_paragraph()
    run = p.add_run(heading)
    run.bold = True
    run.font.size = Pt(12)

    for para_text in text.split("\n\n"):
        p = doc.add_paragraph(para_text)
        pf = p.paragraph_format
        pf.left_indent = Cm(1.5)

page_break()

# ═══════════════════════════════════════════════════════════════════
# 4. SURVEY DATASET OVERVIEW
# ═══════════════════════════════════════════════════════════════════

doc.add_heading("Survey Dataset Overview", level=1)

doc.add_paragraph(
    "The analysis uses the publicly available \"SBLN 2024 Traffic Survey Results\" dataset, comprising "
    "94 spreadsheets commissioned from Intelligent Data Collection Ltd (Project ID-0524-0183) by "
    "Bristol City Council."
)

doc.add_heading("Survey Types", level=2)

add_table(
    ["Survey Type", "Files", "Period", "Coverage"],
    [
        ["ANPR (Origin-Destination)", "6", "11-12 June 2024", "25 cameras across South Bristol"],
        ["ATC (Automatic Traffic Count)", "22", "7-27 June 2024", "22 locations, volume + speed"],
        ["Junction Turning Counts", "22", "11-12 June 2024", "11 junctions"],
        ["Bus Stop Boarding/Alighting", "19", "11-12 June 2024", "Bus passenger counts"],
        ["Queue Length Surveys", "22", "11-12 June 2024", "11 junctions"],
        ["Manual Classified Counts", "4", "8-21 June 2024", "Vehicle classification + cycles"],
    ]
)

doc.add_paragraph(
    "Additional sources include Census 2021 ward-level data (ONS) for Southville and Bedminster wards, "
    "TRICS trip generation rates for urban residential areas, and the official KML file with ANPR camera "
    "GPS coordinates."
)

doc.add_heading("Survey Contractor", level=2)
doc.add_paragraph(
    "Intelligent Data Collection Ltd (Project ID-0524-0183), commissioned by Bristol City Council."
)

doc.add_heading("Data Period", level=2)
doc.add_paragraph(
    "The surveys were conducted between 7 and 27 June 2024. ANPR, JTC, bus, and queue surveys "
    "covered 11-12 June 2024 (two weekdays). ATC counters ran continuously for approximately three weeks. "
    "Manual classified counts covered 8-21 June 2024."
)

page_break()

# ═══════════════════════════════════════════════════════════════════
# 5. ANPR ANALYSIS
# ═══════════════════════════════════════════════════════════════════

doc.add_heading("ANPR Analysis", level=1)

doc.add_heading("Camera Placement", level=2)
doc.add_paragraph(
    "25 ANPR cameras were deployed across South Bristol on 11-12 June 2024 (06:00-19:00). "
    "Critically, none of these cameras were located within the Southville Zone itself. Only two cameras "
    "were positioned on the zone's perimeter roads: Camera 14 (Coronation Road) and Camera 16 (North Street)."
)

add_image("chart_map_anpr.png", Inches(6), "ANPR camera locations (generated map) — none within the Southville Zone")

doc.add_paragraph("Matt's version of the ANPR map, showing the SBLN Perimeter boundary:").runs[0].italic = True
try:
    doc.add_picture('D:/Projects/Bristol/maps/ANPR Cameras Map.png', width=Inches(6))
    doc.add_paragraph("ANPR camera locations (Matt's map) — showing both Southville Zone and SBLN Perimeter").italic = True
except:
    doc.add_paragraph("[Matt's ANPR map not found]")

doc.add_heading("Cross-Zone Detections", level=2)
doc.add_paragraph(
    "The following tables show vehicles detected at both Camera 14 and Camera 16 — the only two cameras "
    "bordering the Southville Zone. These are the maximum possible cut-through vehicles, though even these "
    "may have used perimeter roads."
)

doc.add_paragraph("11 June 2024:").runs[0].bold = True
add_table(
    ["Direction", "<15 min", ">15 min", "Total"],
    [
        ["Cam 14 → Cam 16", "2", "4", "6"],
        ["Cam 16 → Cam 14", "4", "1", "5"],
        ["Day Total", "", "", "11"],
    ]
)

doc.add_paragraph("12 June 2024:").runs[0].bold = True
add_table(
    ["Direction", "<15 min", ">15 min", "Total"],
    [
        ["Cam 14 → Cam 16", "1", "4", "5"],
        ["Cam 16 → Cam 14", "2", "3", "5"],
        ["Day Total", "", "", "10"],
    ]
)

doc.add_paragraph(
    "Total camera traffic on 11 June: Camera 14 recorded 15,092 detections; Camera 16 recorded 7,642. "
    "The cross-zone vehicles represent just 0.07% of Camera 14 traffic and 0.14% of Camera 16 traffic."
)

add_image("chart2_context.png", Inches(6), "Cross-zone detections in context of total traffic")

doc.add_heading("Camera 14 Destination Analysis", level=2)
doc.add_paragraph(
    "Where did vehicles detected at Camera 14 (Coronation Road) go next (within 15 minutes)?"
)
add_table(
    ["Destination", "Camera", "Count", "Share"],
    [
        ["Clift House Rd", "Cam 1", "924", "90%"],
        ["Coronation Rd E", "Cam 7", "75", "~7%"],
        ["North St", "Cam 16", "2", "<1%"],
    ]
)
doc.add_paragraph(
    "The overwhelming majority (90%) of Camera 14 traffic proceeded to Camera 1 at Clift House Road, "
    "indicating travel along Coronation Road — the northern perimeter — not through the Southville Zone."
)
add_image("chart3_cam14_destinations.png", Inches(4), "Camera 14 destination breakdown")

doc.add_heading("Camera 16 Destination Analysis", level=2)
doc.add_paragraph(
    "Where did vehicles detected at Camera 16 (North Street) go next (within 15 minutes)?"
)
add_table(
    ["Destination", "Camera", "Count", "Share"],
    [
        ["Clift House Rd", "Cam 1", "629", "45%"],
        ["East St / A38", "Cam 4", "187", "13%"],
        ["Coronation Rd E", "Cam 7", "112", "8%"],
        ["Coronation Rd", "Cam 14", "4", "<1%"],
    ]
)
doc.add_paragraph(
    "45% of Camera 16 traffic went to Camera 1 (Clift House Road), consistent with travel along "
    "North Street — the southern perimeter — without cutting through the zone."
)
add_image("chart3b_cam16_destinations.png", Inches(4), "Camera 16 destination breakdown")

doc.add_heading("Trip Chain Analysis", level=2)
doc.add_paragraph(
    "Trip chain data tracks vehicles across multiple camera detections to understand journey patterns. "
    "The ANPR data shows that the vast majority of trips detected at Cameras 14 and 16 follow the "
    "perimeter corridors (Coronation Road and North Street) rather than passing through the residential "
    "interior of the Southville Zone."
)

doc.add_heading("Dominant Flows", level=2)
doc.add_paragraph(
    "The dominant traffic flows in the ANPR data are along the major corridors surrounding the "
    "Southville Zone, not through it. The primary flow corridor is Coronation Road (east-west), "
    "with secondary flows along North Street and the A38."
)
add_image("chart5_flows.png", Inches(6), "Dominant traffic flows")
add_image("chart12_flow_map.png", Inches(6), "ANPR flow map")

page_break()

# ═══════════════════════════════════════════════════════════════════
# 6. ATC ANALYSIS
# ═══════════════════════════════════════════════════════════════════

doc.add_heading("ATC Analysis", level=1)

doc.add_paragraph(
    "22 Automatic Traffic Counter (ATC) sites were deployed across South Bristol from 7-27 June 2024, "
    "providing continuous volume and speed data over approximately three weeks. None of the ATC sites "
    "were located within the Southville Zone."
)

add_image("chart_map_atc.png", Inches(6), "ATC site locations (generated map)")

doc.add_paragraph("Matt's version of the ATC map, showing the SBLN Perimeter boundary:").runs[0].italic = True
try:
    doc.add_picture('D:/Projects/Bristol/maps/ATC Sensors Map.png', width=Inches(6))
    doc.add_paragraph("ATC sensor locations (Matt's map) — showing both Southville Zone and SBLN Perimeter").italic = True
except:
    doc.add_paragraph("[Matt's ATC map not found]")

doc.add_heading("Daily Volumes", level=2)
doc.add_paragraph("Key daily volumes (Tuesday-Thursday average):")

add_table(
    ["Site", "Road", "Daily Volume", "AM Peak", "PM Peak"],
    [
        ["12", "Winterstoke Rd (A3029)", "28,602", "-", "2,161"],
        ["8", "Coronation Rd (east)", "23,826", "2,956", "3,318"],
        ["9", "Coronation Rd (west)", "18,894", "2,101", "2,593"],
        ["11", "North St (B3120)", "9,688", "1,363", "1,594"],
        ["7", "Bedminster Parade (A38)", "9,424", "-", "615"],
        ["10", "Luckwell Rd (W-E)", "6,425", "1,074", "1,119"],
        ["15", "Luckwell Rd (S-N)", "5,405", "763", "944"],
        ["14", "Palmyra Rd", "1,958", "426", "349"],
    ]
)

add_image("chart1_volumes.png", Inches(6), "ATC daily traffic volumes")

doc.add_heading("Hourly Profiles", level=2)
doc.add_paragraph(
    "Hourly traffic profiles show classic commuter patterns with AM and PM peaks on weekdays. "
    "The profiles are consistent across all sites, suggesting predictable, locally-driven demand "
    "rather than erratic cut-through behaviour."
)
add_image("chart6_hourly.png", Inches(6), "Hourly traffic profiles")

doc.add_heading("Weekday vs Weekend", level=2)
doc.add_paragraph(
    "Weekend traffic volumes drop compared to weekdays, consistent with commuter-driven demand:"
)
add_table(
    ["Road", "Weekend Drop"],
    [
        ["Luckwell Rd (W-E)", "-16%"],
        ["Luckwell Rd (S-N)", "-9%"],
        ["Palmyra Rd", "-22%"],
    ]
)
doc.add_paragraph(
    "The weekend drops are modest and consistent with residential areas where weekday commuting "
    "accounts for a portion of traffic. A significant cut-through problem would likely show a larger "
    "weekend drop."
)

doc.add_heading("Speed Data", level=2)
doc.add_paragraph(
    "Speed data from nearby residential roads:"
)
add_table(
    ["Road", "Limit", "Mean Speed", "85th %ile", "Exceeding"],
    [
        ["Luckwell Rd (W-E)", "30 mph", "20.7 mph", "~22 mph", "1.8%"],
        ["Palmyra Rd", "20 mph", "17.1 mph", "~22 mph", "21.4%"],
        ["Luckwell Rd (S-N)", "30 mph", "16.6 mph", "~22 mph", "0.4%"],
    ]
)
doc.add_paragraph(
    "Mean speeds are well below posted limits on all measured roads. Palmyra Road has the highest "
    "rate of limit-exceeding vehicles (21.4%), though this is on a 20 mph road where the mean speed "
    "is still only 17.1 mph."
)
add_image("chart4_speeds.png", Inches(6), "Speed data")

page_break()

# ═══════════════════════════════════════════════════════════════════
# 7. JUNCTION TURNING COUNTS
# ═══════════════════════════════════════════════════════════════════

doc.add_heading("Junction Turning Counts", level=1)

doc.add_paragraph(
    "11 junctions were surveyed on 11-12 June 2024. None of the junctions were within the "
    "Southville Zone; only Sites 6 and 11 are on the perimeter."
)

add_image("chart_map_jtc.png", Inches(6), "Junction turning count site locations (generated map)")

doc.add_paragraph("Matt's version of the JTC map, showing the SBLN Perimeter boundary:").runs[0].italic = True
try:
    doc.add_picture('D:/Projects/Bristol/maps/JTC Surveys Map.png', width=Inches(6))
    doc.add_paragraph("JTC survey locations (Matt's map) — showing both Southville Zone and SBLN Perimeter").italic = True
except:
    doc.add_paragraph("[Matt's JTC map not found]")

doc.add_heading("Site 6: Dean Lane / Cannon St / North St Roundabout", level=2)

try:
    doc.add_picture('D:/Projects/Bristol/maps/JTC 6.png', width=Inches(3))
    doc.add_paragraph("Site 6 junction layout (Matt's image)").runs[0].italic = True
except:
    pass

doc.add_paragraph(
    "Cannon Street approaches carry approximately 6,700 vehicles per day, making it a major feeder "
    "from the south. Traffic disperses to Dean Lane (~3,844/day) and westbound North Street (~2,700/day). "
    "Critically, none of the three roads at this junction enter the Southville Zone."
)

doc.add_heading("Site 11: Raleigh Rd / North St Crossroads", level=2)

try:
    doc.add_picture('D:/Projects/Bristol/maps/JTC 11.png', width=Inches(3))
    doc.add_paragraph("Site 11 junction layout (Matt's image)").runs[0].italic = True
except:
    pass
doc.add_paragraph(
    "North Street through-traffic: approximately 3,400 vehicles in each direction over 12 hours. "
    "Around 880 vehicles per day turn between North Street and Raleigh Road, representing less than "
    "10% of the 10,361 total movements at this junction. Raleigh Road does enter the Southville Zone, "
    "but there is no data on how many of these vehicles are through-traffic versus local."
)

add_image("chart13_jtc.png", Inches(6), "Junction turning count analysis")

page_break()

# ═══════════════════════════════════════════════════════════════════
# 8. BUS / PEDESTRIAN & CYCLIST DATA
# ═══════════════════════════════════════════════════════════════════

doc.add_heading("Bus, Pedestrian & Cyclist Data", level=1)

doc.add_heading("Bus Stop Passenger Volumes", level=2)
doc.add_paragraph(
    "19 files cover approximately 10 paired bus stop locations, surveyed on 11-12 June 2024. "
    "Total passenger movements across all sites over the two days: approximately 15,500."
)

add_table(
    ["Site", "Location", "2-Day Passengers"],
    [
        ["5a", "Bedminster Parade NB", "5,380"],
        ["5b", "Bedminster Parade SB", "4,246"],
        ["1a", "North St EB", "1,090"],
        ["3a", "North St EB (mid)", "842"],
        ["1b", "North St WB", "822"],
        ["9b", "West St SB", "746"],
        ["3b", "North St WB (mid)", "538"],
        ["2b", "North St WB (east)", "438"],
        ["2a", "North St EB (east)", "364"],
        ["6a", "St John's Rd SB", "372"],
    ]
)

add_image("chart_map_bus_stops.png", Inches(6), "Bus stop survey locations")
add_image("chart_bus_stops.png", Inches(6), "Bus stop passenger volumes")

doc.add_heading("North Street: Bus vs Car Comparison", level=2)
doc.add_paragraph(
    "North Street bus stops (Sites 1-3) handle approximately 4,100 passenger movements over two days "
    "(~2,050 per day). This indicates significant pedestrian activity along North Street. High bus usage "
    "suggests local residents rely on public transport, not private cars. ATC data shows 9,688 vehicles "
    "per day on North Street — a mix of through-traffic and local access."
)
add_image("chart_bus_vs_cars.png", Inches(6), "North Street bus passengers vs car traffic")

doc.add_heading("Cyclist Data (Manual Classified Counts)", level=2)
doc.add_paragraph(
    "Cycle counts from Winterstoke Road (Site 1, 14 days) show weekday average of approximately 350 "
    "cyclists per day, rising to approximately 500 per day on weekends (higher due to leisure cycling). "
    "Clear commuter patterns are visible in cyclist direction data."
)
doc.add_paragraph(
    "St Luke's Road (Site 12) recorded 228 total cycle/scooter movements on 11 June 2024, with a "
    "strong directional pattern: northbound peaks in AM, southbound peaks in PM (commuters). "
    "E-scooters were counted separately (road vs footpath)."
)
add_image("chart_cyclists.png", Inches(6), "Cyclist counts")

page_break()

# ═══════════════════════════════════════════════════════════════════
# 9. TRIP GENERATION MODEL
# ═══════════════════════════════════════════════════════════════════

doc.add_heading("Trip Generation Model", level=1)

doc.add_heading("Census Data", level=2)
add_table(
    ["Metric", "Southville Ward", "Bedminster Ward", "Combined"],
    [
        ["Population", "12,882", "12,916", "25,798"],
        ["Households", "5,774", "5,780", "11,554"],
        ["No car", "32.2%", "~28%", "~30%"],
        ["Car-owning", "67.8%", "~72%", "~70%"],
    ]
)

doc.add_heading("TRICS Trip Generation Rates", level=2)
doc.add_paragraph(
    "Standard TRICS rates for urban residential areas: 5.0-5.7 vehicle trips per dwelling per day "
    "(12-hour weekday). Applied to an estimated Luckwell Road catchment of approximately 1,800 dwellings:"
)
doc.add_paragraph(
    "Mid-range estimate: ~4,350 vehicles/day via Luckwell Road, which accounts for 68% of the "
    "observed 6,425 daily volume. When schools, deliveries, and visitors are included, approximately "
    "94% of the observed traffic can be explained by local demand."
)

add_image("chart7_trip_generation.png", Inches(6), "Trip generation model")
add_image("chart8_car_ownership.png", Inches(5), "Car ownership by ward")
add_image("chart9_waterfall.png", Inches(6), "Traffic demand waterfall — local sources explaining observed volumes")

page_break()

# ═══════════════════════════════════════════════════════════════════
# 10. COMPREHENSIVE ANALYSIS MAP
# ═══════════════════════════════════════════════════════════════════

doc.add_heading("Comprehensive Analysis Map", level=1)

doc.add_paragraph(
    "The following map combines all survey sensor locations and key findings into a single view."
)

add_image("chart_analysis_map.png", Inches(6), "Comprehensive analysis map — all findings")

doc.add_paragraph()

add_image("chart_map_all_sensors.png", Inches(6), "All sensor locations combined")

page_break()

# ═══════════════════════════════════════════════════════════════════
# 11. METHODOLOGICAL LIMITATIONS
# ═══════════════════════════════════════════════════════════════════

doc.add_heading("Methodological Limitations", level=1)

doc.add_heading("Cut-Through Evidence Gap", level=2)
doc.add_paragraph(
    "The fundamental limitation is the absence of any ANPR cameras within the Southville Zone. "
    "Without interior cameras, it is impossible to determine whether vehicles detected at both "
    "Camera 14 and Camera 16 actually drove through the zone or used perimeter roads. "
    "The 10-11 cross-zone vehicles per day therefore represent an upper bound, not a confirmed count."
)

doc.add_heading("Volume vs Route Data", level=2)
doc.add_paragraph(
    "ATC counters measure traffic volume but cannot determine origin, destination, or route. "
    "High volumes on perimeter roads do not indicate cut-through traffic. The ATC data shows "
    "how many vehicles use a road, not why they are there."
)

doc.add_heading("Sample Size", level=2)
doc.add_paragraph(
    "ANPR and JTC data cover only two survey days (11-12 June 2024). While ATC data covers "
    "three weeks, the origin-destination information critical to the cut-through question is "
    "limited to this two-day sample."
)

doc.add_heading("Camera Position Discrepancy", level=2)
doc.add_paragraph(
    "The official KML file provides GPS coordinates for all 25 ANPR cameras. The camera positions "
    "confirm that no cameras were placed within the Southville Zone's residential streets — they "
    "are all on boundary roads or in other parts of South Bristol entirely."
)

doc.add_heading("2025 Surveys", level=2)
doc.add_paragraph(
    "At the time of writing, no additional surveys have been published that address the evidence gap. "
    "Any future surveys would need interior ANPR coverage to meaningfully assess cut-through traffic."
)

page_break()

# ═══════════════════════════════════════════════════════════════════
# 12. CONCLUSIONS
# ═══════════════════════════════════════════════════════════════════

doc.add_heading("Conclusions", level=1)

conclusions = [
    "The ANPR data shows only 10-11 vehicles per day detected at both Camera 14 (Coronation Road) "
    "and Camera 16 (North Street). This represents 0.07-0.14% of traffic at those cameras — far from "
    "the \"high numbers\" claimed by Bristol City Council.",

    "Even these 10-11 vehicles cannot be confirmed as having driven through the Southville Zone, "
    "because there are no ANPR cameras within the zone to track their route. They may have used "
    "perimeter roads.",

    "The dominant traffic flows at Cameras 14 and 16 are along the perimeter corridors: 90% of "
    "Camera 14 traffic goes to Camera 1 (Clift House Road) via Coronation Road, and 45% of Camera 16 "
    "traffic also goes to Camera 1 via North Street.",

    "ATC data shows traffic volumes consistent with locally-generated demand. Trip generation modelling "
    "using Census 2021 household data and TRICS rates explains approximately 94% of observed traffic on "
    "Luckwell Road through local sources (residents, schools, deliveries, visitors).",

    "No survey instruments — ANPR, ATC, JTC, or MCC — were placed within the Southville Zone's "
    "residential streets. The entire survey programme measured perimeter and boundary roads only.",

    "The claim that \"traffic data show that high numbers of vehicles from outside the area use "
    "residential streets in Southville as a cut-through\" is not supported by the available survey data. "
    "The data required to prove or disprove this claim (interior ANPR coverage) was never collected.",
]

for i, c in enumerate(conclusions, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. ")
    run.bold = True
    p.add_run(c)

page_break()

# ═══════════════════════════════════════════════════════════════════
# 13. MATT'S FEEDBACK LOG
# ═══════════════════════════════════════════════════════════════════

doc.add_heading("Matt's Feedback Log", level=1)

doc.add_paragraph(
    "The following table documents all feedback items from Matt Sanders' reviews of the analysis "
    "and visualisations, with their resolution status."
)

feedback_items = [
    ("Add ANPR camera map with zone boundary", "DONE", "chart_map_anpr.png created"),
    ("Show cross-zone numbers in context", "DONE", "chart2_context.png — percentage bars"),
    ("Camera 14 destination breakdown chart", "DONE", "chart3_cam14_destinations.png"),
    ("Camera 16 destination breakdown chart", "DONE", "chart3b_cam16_destinations.png"),
    ("Dominant flow visualisation", "DONE", "chart5_flows.png"),
    ("Flow map with arrows", "DONE", "chart12_flow_map.png"),
    ("ATC site map", "DONE", "chart_map_atc.png"),
    ("Daily volumes bar chart", "DONE", "chart1_volumes.png"),
    ("Hourly profile chart", "DONE", "chart6_hourly.png"),
    ("Speed data chart", "DONE", "chart4_speeds.png"),
    ("JTC site map", "DONE", "chart_map_jtc.png"),
    ("JTC analysis chart", "DONE", "chart13_jtc.png"),
    ("Trip generation model chart", "DONE", "chart7_trip_generation.png"),
    ("Car ownership chart", "DONE", "chart8_car_ownership.png"),
    ("Waterfall chart for demand explanation", "DONE", "chart9_waterfall.png"),
    ("Comprehensive analysis map", "DONE", "chart_analysis_map.png"),
    ("All sensors combined map", "DONE", "chart_map_all_sensors.png"),
    ("Bus stop map", "DONE", "chart_map_bus_stops.png"),
    ("Bus stop passenger volumes chart", "DONE", "chart_bus_stops.png"),
    ("Bus vs cars comparison chart", "DONE", "chart_bus_vs_cars.png"),
    ("Cyclist data chart", "DONE", "chart_cyclists.png"),
    ("Include full WECA statement text", "DONE", "Section 3 of this document"),
    ("Add data file inventory appendix", "DONE", "Appendix A"),
]

add_table(
    ["#", "Item", "Status", "Notes"],
    [[str(i+1), item, status, notes] for i, (item, status, notes) in enumerate(feedback_items)]
)

page_break()

# ═══════════════════════════════════════════════════════════════════
# 14. APPENDIX: DATA FILE INVENTORY
# ═══════════════════════════════════════════════════════════════════

doc.add_heading("Appendix A: Data File Inventory", level=1)

doc.add_paragraph(
    "Complete list of all files in the SBLN 2024 Traffic Survey Results dataset, organised by survey type."
)

doc.add_heading("ANPR Files (6)", level=2)
anpr_files = [
    "ANPR_OD_Matrix_11June2024.xlsx",
    "ANPR_OD_Matrix_12June2024.xlsx",
    "ANPR_Sample_Rates_11June2024.xlsx",
    "ANPR_Sample_Rates_12June2024.xlsx",
    "ANPR_Trip_Chains_11June2024.xlsx",
    "ANPR_Trip_Chains_12June2024.xlsx",
]
for f in anpr_files:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading("ATC Files (22)", level=2)
for i in range(1, 23):
    doc.add_paragraph(f"ATC_Site{i:02d}_7to27June2024.xlsx", style='List Bullet')

doc.add_heading("Junction Turning Count Files (22)", level=2)
for i in range(1, 12):
    doc.add_paragraph(f"JTC_Site{i:02d}_11June2024.xlsx", style='List Bullet')
    doc.add_paragraph(f"JTC_Site{i:02d}_12June2024.xlsx", style='List Bullet')

doc.add_heading("Bus Stop Boarding/Alighting Files (19)", level=2)
bus_sites = ["1a", "1b", "2a", "2b", "3a", "3b", "4a", "5a", "5b", "6a",
             "6b", "7a", "7b", "8a", "8b", "9a", "9b", "10a", "10b"]
for s in bus_sites:
    doc.add_paragraph(f"Bus_Site{s}_11to12June2024.xlsx", style='List Bullet')

doc.add_heading("Queue Length Survey Files (22)", level=2)
for i in range(1, 12):
    doc.add_paragraph(f"Queue_Site{i:02d}_11June2024.xlsx", style='List Bullet')
    doc.add_paragraph(f"Queue_Site{i:02d}_12June2024.xlsx", style='List Bullet')

doc.add_heading("Manual Classified Count Files (4)", level=2)
mcc_files = [
    "MCC_Site01_LinkCount_8to21June2024.xlsx",
    "MCC_Site02_LinkCount_8to21June2024.xlsx",
    "MCC_Site03_LinkCount_11to12June2024.xlsx",
    "MCC_Site12_CycleScooter_11June2024.xlsx",
]
for f in mcc_files:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Total: 95 files (94 spreadsheets + 1 KML coordinate file)")
run.bold = True

# ── Save ──
doc.save(OUTPUT)
print(f"Document saved to: {OUTPUT}")
